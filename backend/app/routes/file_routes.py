from flask import Blueprint, request, jsonify
from ..services.file_service import FileService
from ..lib.supabase import supabase
from ..utils.logger import Logger
import traceback

file_bp = Blueprint('file', __name__)
logger = Logger()
service = FileService()

@file_bp.route('/upload', methods=['POST'])
def upload_file():
    logger.log_with_timestamp('FILE_UPLOAD', 'Bắt đầu xử lý request upload file')
    
    # Log request details
    logger.log_with_timestamp('FILE_UPLOAD', f'Request form data: {list(request.form.keys())}')
    logger.log_with_timestamp('FILE_UPLOAD', f'Request files: {list(request.files.keys()) if request.files else "No files"}')
    
    # Expect multipart/form-data with 'file', 'user_id', and optional 'space_id'
    user_id = request.form.get('user_id')
    space_id = request.form.get('space_id')  # Optional space_id
    
    if not user_id:
        logger.log_with_timestamp('FILE_UPLOAD_ERROR', 'Missing user_id in request form data')
        return jsonify({'error': 'Missing user_id'}), 400

    logger.log_with_timestamp('FILE_UPLOAD', f'User ID: {user_id}')

    if 'file' not in request.files:
        logger.log_with_timestamp('FILE_UPLOAD_ERROR', 'No file found in request files')
        return jsonify({'error': 'No file provided'}), 400
        
    file = request.files['file']
    if file.filename == '':
        logger.log_with_timestamp('FILE_UPLOAD_ERROR', 'Empty filename')
        return jsonify({'error': 'Empty filename'}), 400
        
    logger.log_with_timestamp('FILE_UPLOAD', f'File received: {file.filename}, Content-Type: {file.content_type}, Size: {file.content_length or "unknown"} bytes')

    try:
        # extract text content (assume text extraction earlier done by caller)
        from ..services.file_service import FileService
        
        # Process different file types
        if file.content_type == 'application/pdf':
            try:
                logger.log_with_timestamp('FILE_UPLOAD', 'Processing PDF file...')
                # For PDF files, we need to use a PDF parser
                import PyPDF2
                import io
                
                # Create a file-like object from the uploaded file
                pdf_file = io.BytesIO(file.read())
                
                # Create PDF reader
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                # Extract text from all pages
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                    
                logger.log_with_timestamp('FILE_UPLOAD', f'Successfully extracted {len(text_content)} characters from PDF')
                content = text_content
            except ImportError:
                logger.log_with_timestamp('FILE_UPLOAD_ERROR', 'PyPDF2 library not installed, trying fallback method')
                # Fallback to binary read if PyPDF2 is not available
                file.seek(0)  # Reset file pointer to beginning
                content = file.read().decode('utf-8', errors='ignore')
        elif file.content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
            try:
                logger.log_with_timestamp('FILE_UPLOAD', 'Processing Word document...')
                from docx import Document
                import io
                import zipfile
                
                # Read file content first
                file_content = file.read()
                logger.log_with_timestamp('FILE_UPLOAD', f'Read {len(file_content)} bytes from uploaded file')
                
                # Create a file-like object from the uploaded file
                if file.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                    # DOCX file - verify it's a valid ZIP file first
                    docx_file = io.BytesIO(file_content)
                    
                    # Try to verify it's a valid ZIP file and process with python-docx
                    try:
                        # Try to open with python-docx directly first
                        doc = Document(docx_file)
                        
                        # Extract text from all paragraphs
                        text_content = ""
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():  # Only add non-empty paragraphs
                                text_content += paragraph.text + "\n"
                        
                        # Extract text from tables if any
                        for table in doc.tables:
                            for row in table.rows:
                                row_text = ""
                                for cell in row.cells:
                                    if cell.text.strip():
                                        row_text += cell.text.strip() + " | "
                                if row_text:
                                    text_content += row_text.rstrip(" | ") + "\n"
                        
                        # Remove excessive whitespace
                        text_content = '\n'.join(line.strip() for line in text_content.split('\n') if line.strip())
                        
                        if text_content:
                            logger.log_with_timestamp('FILE_UPLOAD', f'Successfully extracted {len(text_content)} characters from DOCX')
                            content = text_content
                        else:
                            logger.log_with_timestamp('FILE_UPLOAD', 'No text content found in DOCX, using fallback')
                            raise Exception("No readable text content found")
                            
                    except Exception as docx_error:
                        logger.log_with_timestamp('FILE_UPLOAD_ERROR', f'python-docx processing failed: {str(docx_error)}, trying ZIP verification')
                        
                        # Try ZIP verification as fallback
                        try:
                            docx_file.seek(0)
                            with zipfile.ZipFile(docx_file, 'r') as zip_file:
                                # Check if it has the typical DOCX structure
                                if 'word/document.xml' not in zip_file.namelist():
                                    raise Exception("Not a valid DOCX file structure")
                            
                            # If ZIP is valid, try python-docx again
                            docx_file.seek(0)
                            doc = Document(docx_file)
                            
                            # Extract text from all paragraphs
                            text_content = ""
                            for paragraph in doc.paragraphs:
                                if paragraph.text.strip():  # Only add non-empty paragraphs
                                    text_content += paragraph.text + "\n"
                            
                            # Extract text from tables if any
                            for table in doc.tables:
                                for row in table.rows:
                                    row_text = ""
                                    for cell in row.cells:
                                        if cell.text.strip():
                                            row_text += cell.text.strip() + " | "
                                    if row_text:
                                        text_content += row_text.rstrip(" | ") + "\n"
                            
                            # Remove excessive whitespace
                            text_content = '\n'.join(line.strip() for line in text_content.split('\n') if line.strip())
                            
                            if text_content:
                                logger.log_with_timestamp('FILE_UPLOAD', f'Successfully extracted {len(text_content)} characters from DOCX on retry')
                                content = text_content
                            else:
                                logger.log_with_timestamp('FILE_UPLOAD', 'No text content found in DOCX on retry, using fallback')
                                raise Exception("No readable text content found on retry")
                                
                        except zipfile.BadZipFile:
                            logger.log_with_timestamp('FILE_UPLOAD_ERROR', 'File is not a valid ZIP/DOCX file, using fallback')
                            raise Exception("Invalid DOCX format")
                        except Exception as retry_error:
                            logger.log_with_timestamp('FILE_UPLOAD_ERROR', f'Retry also failed: {str(retry_error)}, using fallback')
                            raise Exception("DOCX processing failed completely")
                        
                else:
                    # DOC file - use alternative extraction method
                    logger.log_with_timestamp('FILE_UPLOAD', 'DOC file detected, trying text extraction')
                    try:
                        # Try to extract readable text from DOC file
                        import re
                        # Convert bytes to string and extract readable text
                        raw_text = file_content.decode('utf-8', errors='ignore')
                        # Remove control characters and keep only printable text
                        text_content = re.sub(r'[^\x20-\x7E\n\r\t]', ' ', raw_text)
                        # Clean up multiple spaces and empty lines
                        text_content = re.sub(r'\s+', ' ', text_content)
                        text_content = re.sub(r'\n\s*\n', '\n', text_content)
                        
                        if len(text_content.strip()) > 50:  # If we got reasonable amount of text
                            logger.log_with_timestamp('FILE_UPLOAD', f'Extracted {len(text_content)} characters from DOC file')
                            content = text_content
                        else:
                            raise Exception("Insufficient readable content from DOC file")
                    except Exception as doc_error:
                        logger.log_with_timestamp('FILE_UPLOAD_ERROR', f'DOC processing failed: {str(doc_error)}, using raw fallback')
                        content = file_content.decode('utf-8', errors='ignore')
                    
            except ImportError:
                logger.log_with_timestamp('FILE_UPLOAD_ERROR', 'python-docx library not available, trying fallback method')
                # Fallback to binary read if python-docx is not available
                file.seek(0)  # Reset file pointer to beginning
                content = file.read().decode('utf-8', errors='ignore')
            except Exception as e:
                logger.log_with_timestamp('FILE_UPLOAD_ERROR', f'Error processing Word document: {str(e)}, trying fallback method')
                # Fallback to binary read if document processing fails
                file.seek(0)  # Reset file pointer to beginning
                content = file.read().decode('utf-8', errors='ignore')
        else:
            # For text files, just read as text
            content = file.read().decode('utf-8', errors='ignore')
            logger.log_with_timestamp('FILE_UPLOAD', f'Extracted {len(content)} characters from text file')
        
        # Log content sample for debugging
        content_sample = content[:200] + "..." if len(content) > 200 else content
        logger.log_with_timestamp('FILE_UPLOAD', f'Content sample: {content_sample}')
        
        # Reset file pointer for FileService
        file.seek(0)
        
        # process and save - Đã chuyển về gọi hàm đồng bộ
        logger.log_with_timestamp('FILE_UPLOAD', 'Saving to Supabase...')
        file_id = service.save_file_and_chunks_to_supabase(user_id, file, content, space_id)
        logger.log_with_timestamp('FILE_UPLOAD', f'Successfully saved file with ID: {file_id}')
        return jsonify({'success': True, 'file_id': file_id, 'filename': file.filename})
    except Exception as e:
        error_trace = traceback.format_exc()
        logger.log_with_timestamp('FILE_UPLOAD_ERROR', f'Error: {str(e)}\nTraceback: {error_trace}')
        return jsonify({'error': str(e)}), 500

@file_bp.route('/list', methods=['GET'])
def list_files():
    user_id = request.args.get('user_id')
    if not user_id:
        return jsonify({'error': 'Missing user_id'}), 400
    try:
        res = supabase.table('user_files').select('id, filename, created_at').eq('user_id', user_id).order('created_at', desc=True).execute()
        data = res.data or []
        return jsonify({'files': data})
    except Exception as e:
        logger.log_with_timestamp('FILE_ROUTE_ERROR', str(e))
        return jsonify({'error': str(e)}), 500

@file_bp.route('/<uuid:file_id>', methods=['DELETE'])
def delete_file(file_id):
    user_id = request.args.get('user_id')
    if not user_id:
        return jsonify({'error': 'Missing user_id'}), 400
    try:
        # delete metadata (cascade deletes chunks)
        res = supabase.table('user_files').delete().eq('id', str(file_id)).eq('user_id', user_id).execute()
        # Check for errors in the new Supabase client format
        if hasattr(res, 'error') and res.error:
            raise Exception(res.error.message)
        elif not res.data:
            raise Exception("File not found or already deleted")
        return jsonify({'success': True})
    except Exception as e:
        logger.log_with_timestamp('FILE_ROUTE_ERROR', str(e))
        return jsonify({'error': str(e)}), 500